import streamlit as st
import google.generativeai as genai
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import PIL.Image
import PyPDF2
import re

# --- 1. CONFIGURATION ---
st.set_page_config(
    page_title="Magis AI - Ignatian Pedagogy",
    page_icon="https://i.imgur.com/UUCgyfV.png",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- 2. TRANSLATION DICTIONARY (LENGKAP) ---
LANGUAGES = {
    "Bahasa Indonesia": {
        "sidebar_header": "Pengaturan",
        "lang_label": "🌐 Bahasa / Language",
        "api_key_label": "🔑 Kunci Akses (API Key)",
        "status_connected": "✅ Sistem {provider} Terhubung",
        "status_waiting": "⚠️ Menunggu API Key",
        "division_label": "📌 Pilih Divisi / Mode",
        # Chat Bebas dikembalikan
        "divisions": ["Akademik (Pedagogi)", "Pastoral & Diskresi", "Manajemen Sekolah", "Chat Bebas (General)"],
        
        # Akademik
        "academic_header": "#### 🎓 Konfigurasi Akademik",
        "class_expander": "📚 Kelas & Materi",
        "grade_label": "Jenjang Kelas",
        "grades": ["7 SMP", "8 SMP", "9 SMP", "10 SMA (Fase E)", "11 SMA (Fase F)", "12 SMA (Fase F)"],
        "subject_label": "Mata Pelajaran",
        "subject_ph": "Misal: Sejarah Indonesia",
        "topic_label": "Kompetensi Dasar (KD) / CP",
        "topic_ph": "Paste CP/Tujuan Pembelajaran...",
        "params_expander": "🧠 Parameter Soal & Tugas",
        "bloom_label": "Level Kognitif (Bloom)",
        "bloom_opts": ["C1 (Mengingat)", "C2 (Memahami)", "C3 (Menerapkan)", "C4 (Menganalisis)", "C5 (Mengevaluasi)", "C6 (Mencipta)"],
        "difficulty_label": "Tingkat Kesulitan",
        "difficulty_opts": ["Mudah", "Sedang", "HOTS (Sulit)", "Olimpiade"],
        "style_expander": "🎨 Gaya & Pendekatan Ignasian",
        "style_label": "Gaya Bahasa",
        "style_opts": ["Formal Akademis", "Sokratik (Bertanya Balik)", "Storytelling (Naratif)", "Simpel & Lugas"],
        "ipp_label": "Fokus IPP",
        
        # Pastoral
        "pastoral_header": "#### 🕊️ Pendampingan Pastoral",
        "counsel_expander": "❤️ Konteks Konseling",
        "subject_label": "Subjek",
        "subjects": ["Siswa", "Guru/Karyawan", "Orang Tua", "Alumni"],
        "issue_label": "Isu Utama",
        "issue_opts": ["Akademik", "Keluarga", "Pencarian Jati Diri", "Keputusan Besar (Diskresi)", "Kejenuhan/Burnout", "Lainnya (Tulis Sendiri)..."],
        "issue_custom_ph": "Misal: Konflik dengan teman sebaya...",
        "method_label": "Metode Pendampingan",
        "methods": ["Mendengarkan (Listening)", "Diskresi (Pembedaan Roh)", "Examen (Refleksi Harian)"],
        
        # Manajemen
        "management_header": "#### 💼 Manajemen Sekolah",
        "doc_type_label": "Jenis Dokumen",
        "doc_types": ["Surat Resmi", "Proposal Kegiatan", "Pidato/Sambutan", "Email Internal"],
        "tone_label": "Nada Bicara",
        "tones": ["Tegas & Formal", "Persuasif", "Apresiatif", "Instruktif"],
        "topic_mgmt_label": "Topik/Acara",
        "topic_mgmt_ph": "Misal: Hari Guru",

        # General / Chat Bebas
        "general_header": "#### 💬 Diskusi Bebas (General)",
        "general_help": "Mode ini untuk brainstorming ide, bertanya hal umum, atau diskusi santai tanpa template khusus.",
        
        # Aksesibilitas
        "access_header": "👁️ Tampilan & Aksesibilitas",
        "font_size_label": "Ukuran Teks",
        
        # UI & Footer
        "reset_btn": "🔄 Reset Sesi Chat",
        "magis_title": "MAGIS AI",
        "magis_tagline": "Mitra Diskresi Guru Ignasian",
        "mode_active": "Mode Aktif",
        "upload_expander": "📂 Upload Dokumen & Materi Referensi",
        "upload_help": "Upload RPP, E-Book, atau Gambar Soal untuk dianalisis AI.",
        "file_label": "Pilih file (PDF, Docx, TXT)",
        "img_label": "Upload Gambar (Jika perlu)",
        "success_msg": "📚 {count} dokumen berhasil dipelajari.",
        "workspace_header": "### ✍️ Area Kerja",
        "instruction_label": "Instruksi / Pertanyaan:",
        "submit_btn": "🚀 KIRIM PERINTAH",
        "tip_msg": "💡 *Tip: Semakin detail instruksi, semakin tajam hasil analisis Ignasian.*",
        "spinner_msg": "✨ Sedang meracik materi dengan perspektif Ignasian...",
        "export_header": "### 📥 Ekspor Hasil",
        "download_label": "Download Dokumen Word (.docx)",
        "system_lang_instruction": "Jawablah dalam Bahasa Indonesia yang akademis dan reflektif.",
        "footer_design": "Design by: Albertus Henny Setyawan"
    },
    "English": {
        "sidebar_header": "Settings",
        "lang_label": "🌐 Language / Bahasa",
        "api_key_label": "🔑 API Key",
        "status_connected": "✅ {provider} System Connected",
        "status_waiting": "⚠️ Waiting for API Key",
        "division_label": "📌 Select Service Division",
        "divisions": ["Academic (Pedagogy)", "Pastoral & Discernment", "School Management", "General Chat (Free)"],
        
        # Academic
        "academic_header": "#### 🎓 Academic Configuration",
        "class_expander": "📚 Class & Material",
        "grade_label": "Grade Level",
        "grades": ["7 Junior High", "8 Junior High", "9 Junior High", "10 Senior High (Phase E)", "11 Senior High (Phase F)", "12 Senior High (Phase F)"],
        "subject_label": "Subject",
        "subject_ph": "E.g., Indonesian History",
        "topic_label": "Competency / Learning Objectives",
        "topic_ph": "Paste Learning Objectives here...",
        "params_expander": "🧠 Questions & Tasks Parameters",
        "bloom_label": "Cognitive Level (Bloom)",
        "bloom_opts": ["C1 (Remembering)", "C2 (Understanding)", "C3 (Applying)", "C4 (Analyzing)", "C5 (Evaluating)", "C6 (Creating)"],
        "difficulty_label": "Difficulty Level",
        "difficulty_opts": ["Easy", "Medium", "HOTS (Hard)", "Olympiad"],
        "style_expander": "🎨 Ignatian Style & Approach",
        "style_label": "Language Style",
        "style_opts": ["Formal Academic", "Socratic (Questioning)", "Storytelling (Narrative)", "Simple & Direct"],
        "ipp_label": "IPP Focus",
        
        # Pastoral
        "pastoral_header": "#### 🕊️ Pastoral Care",
        "counsel_expander": "❤️ Counseling Context",
        "subject_label": "Subject",
        "subjects": ["Student", "Teacher/Staff", "Parent", "Alumni"],
        "issue_label": "Main Issue",
        "issue_opts": ["Academic", "Family", "Identity Search", "Major Decision (Discernment)", "Burnout", "Other (Type below)..."],
        "issue_custom_ph": "E.g., Conflict with peers...",
        "method_label": "Approach Method",
        "methods": ["Listening", "Discernment", "Examen (Daily Reflection)"],
        
        # Management
        "management_header": "#### 💼 School Management",
        "doc_type_label": "Document Type",
        "doc_types": ["Official Letter", "Event Proposal", "Speech/Remarks", "Internal Email"],
        "tone_label": "Tone of Voice",
        "tones": ["Firm & Formal", "Persuasive", "Appreciative", "Instructive"],
        "topic_mgmt_label": "Topic/Event",
        "topic_mgmt_ph": "E.g., Teacher's Day",

        # General
        "general_header": "#### 💬 General Discussion",
        "general_help": "Use this mode for brainstorming, general questions, or casual discussion without specific templates.",
        
        # Accessibility
        "access_header": "👁️ Appearance & Accessibility",
        "font_size_label": "Text Size",

        # UI
        "reset_btn": "🔄 Reset Chat Session",
        "magis_title": "MAGIS AI",
        "magis_tagline": "Ignatian Teacher's Discernment Partner",
        "mode_active": "Active Mode",
        "upload_expander": "📂 Upload Documents & Reference",
        "upload_help": "Upload Lesson Plans, E-Books, or Image Questions for AI analysis.",
        "file_label": "Select files (PDF, Docx, TXT)",
        "img_label": "Upload Image (If needed)",
        "success_msg": "📚 {count} documents successfully processed.",
        "workspace_header": "### ✍️ Workspace",
        "instruction_label": "Instruction / Question:",
        "submit_btn": "🚀 SEND COMMAND",
        "tip_msg": "💡 *Tip: The more detailed the instruction, the sharper the Ignatian analysis.*",
        "spinner_msg": "✨ Crafting material with Ignatian perspective...",
        "export_header": "### 📥 Export Results",
        "download_label": "Download Word Document (.docx)",
        "system_lang_instruction": "Answer in English using an academic and reflective tone suitable for a Jesuit institution.",
        "footer_design": "Design by: Albertus Henny Setyawan"
    }
}

# --- 3. CSS MODERN & DYNAMIC THEME (DENGAN FONT SIZE) ---
def inject_css(base_size_px):
    # Hitung rasio font untuk elemen lain
    h1_size = int(base_size_px * 2.5)
    h2_size = int(base_size_px * 2.0)
    h3_size = int(base_size_px * 1.5)
    small_size = int(base_size_px * 0.85)
    
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');
    
    /* GLOBAL THEME */
    html, body, [class*="css"] {{
        font-family: 'Plus Jakarta Sans', sans-serif;
    }}
    
    .stApp {{ 
        background-color: #F8F9FB !important; 
        color: #1E293B !important; 
    }}

    /* DYNAMIC FONT SIZES */
    p, .stMarkdown, .stText, .stTextArea textarea, .stSelectbox, .stTextInput {{
        font-size: {base_size_px}px !important;
    }}
    
    /* SIDEBAR STYLING */
    section[data-testid="stSidebar"] {{ 
        background-color: #FFFFFF !important; 
        border-right: 1px solid #E2E8F0;
        box-shadow: 4px 0 24px rgba(0,0,0,0.02);
    }}
    
    /* TYPOGRAPHY */
    h1 {{ color: #1B365D !important; letter-spacing: -0.5px; font-size: {h1_size}px !important; }}
    h2 {{ color: #1B365D !important; letter-spacing: -0.5px; font-size: {h2_size}px !important; }}
    h3 {{ color: #1B365D !important; letter-spacing: -0.5px; font-size: {h3_size}px !important; }}
    
    /* HEADER JUDUL (GRADIENT TEXT) */
    .magis-title {{
        font-weight: 800; 
        font-size: {h1_size}px; 
        background: linear-gradient(135deg, #1B365D 0%, #B8860B 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        margin-bottom: 0px;
        line-height: 1.1;
        letter-spacing: -1px;
    }}
    
    .magis-tagline {{
        font-size: {base_size_px + 2}px;
        font-weight: 600;
        font-style: italic;
        color: #576F8E; 
        margin-bottom: 15px;
        border-left: 3px solid #DAA520;
        padding-left: 10px;
    }}

    .magis-badge {{
        display: inline-block;
        background-color: #E0F2FE;
        color: #0284C7;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: {small_size}px;
        font-weight: 700;
        margin-bottom: 20px;
    }}
    
    /* CHAT BUBBLES */
    .bubble-user {{
        background: linear-gradient(135deg, #1B365D 0%, #2D4F85 100%);
        color: white; 
        padding: 20px; 
        border-radius: 20px 20px 4px 20px; 
        margin-left: auto; max-width: 85%;
        box-shadow: 0 10px 15px -3px rgba(27, 54, 93, 0.2);
        font-size: {base_size_px}px;
        line-height: 1.6;
    }}
    .bubble-ai {{
        background-color: #FFFFFF; 
        color: #334155; 
        border: 1px solid #F1F5F9; 
        border-left: 5px solid #DAA520;
        padding: 24px; 
        border-radius: 4px 20px 20px 20px; 
        margin-right: auto; max-width: 95%;
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);
        font-size: {base_size_px}px;
        line-height: 1.6;
    }}
    
    /* INPUT AREA & FORM STYLING */
    .stTextArea textarea {{
        background-color: #FFFFFF !important;
        border: 2px solid #E2E8F0 !important;
        border-radius: 12px !important;
        padding: 15px !important;
        font-size: {base_size_px}px !important;
        transition: all 0.3s ease;
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.02);
    }}
    .stTextArea textarea:focus {{
        border-color: #1B365D !important;
        box-shadow: 0 0 0 3px rgba(27, 54, 93, 0.1) !important;
    }}
    
    /* BUTTON STYLING */
    div[data-testid="stForm"] button {{
        background: linear-gradient(90deg, #1B365D 0%, #162B4A 100%);
        color: white;
        font-weight: 700;
        border-radius: 12px;
        padding: 10px 0;
        border: none;
        transition: transform 0.2s, box-shadow 0.2s;
        box-shadow: 0 4px 6px rgba(27, 54, 93, 0.2);
        font-size: {base_size_px}px !important;
    }}
    div[data-testid="stForm"] button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 10px 15px rgba(27, 54, 93, 0.3);
    }}
    
    /* STATUS INDICATORS */
    .status-ok {{ color: #059669; font-weight: bold; font-size: {small_size}px; border: 1px solid #059669; padding: 8px; border-radius: 8px; background: #ECFDF5; display: flex; align-items: center; gap: 5px;}}
    .status-err {{ color: #DC2626; font-weight: bold; font-size: {small_size}px; border: 1px solid #DC2626; padding: 8px; border-radius: 8px; background: #FEF2F2; display: flex; align-items: center; gap: 5px;}}
    
    /* FOOTER */
    .sidebar-footer {{
        text-align: center;
        margin-top: 30px;
        padding-top: 20px;
        border-top: 1px dashed #CBD5E1;
        color: #64748B;
        font-size: {small_size}px;
        line-height: 1.5;
    }}
    </style>
    """, unsafe_allow_html=True)

# --- 4. IGNATIAN DNA (SYSTEM PROMPT) ---
def get_system_prompt(language_instruction):
    return f"""
PERAN: 'Magis AI', asisten pedagogi dan pendamping spiritual khas Kolese Jesuit (Ignasian).

DNA SPIRITUAL & FILOSOFI (WAJIB DIINTEGRASIKAN):
1. IPP (Ignatian Pedagogical Paradigm): Context, Experience, Reflection, Action, Evaluation.
2. Cura Personalis: Perhatian personal pada keunikan tiap pribadi.
3. Universal Apostolic Preferences (UAP): Menunjukkan jalan menuju Tuhan, Berjalan bersama yang tersingkir, Menemani kaum muda, Merawat rumah bersama.
4. Ignatian Leadership & Values: 4C (Competence, Conscience, Compassion, Commitment).
5. Diskresi (Pembedaan Roh) & Spiritualitas: Membedakan gerak roh baik (Consolation) dan roh jahat (Desolation), AMDG.
6.  **IPP (Ignatian Pedagogical Paradigm):**
    - **Context:** Memahami dunia nyata siswa, keluarga, dan budaya.
    - **Experience:** Melibatkan rasa, akal budi, dan imajinasi (bukan hafalan semata).
    - **Reflection:** Menggali makna, nilai, dan kebenaran dari pengalaman.
    - **Action:** Dorongan untuk bertindak melayani sesama (Magis).
    - **Evaluation:** Menilai perkembangan karakter dan kompetensi secara utuh.

7.  **Cura Personalis:**
    - Perhatian personal pada keunikan tiap pribadi.
    - Mendengarkan dengan hati, memanusiakan, dan tidak menghakimi.

8.  **Universal Apostolic Preferences (UAP):**
    - Menunjukkan jalan menuju Tuhan (Latihan Rohani & Diskresi).
    - Berjalan bersama yang tersingkir (Solidaritas & Keadilan).
    - Menemani kaum muda menuju masa depan penuh harapan.
    - Merawat rumah kita bersama (Kesadaran Ekologis).

9.  **Ignatian Leadership & Values:**
    - **Self-Awareness:** Mengenal diri (kekuatan & kelemahan) dalam terang Ilahi.
    - **Ingenuity:** Cerdik dan inovatif dalam menghadapi tantangan zaman.
    - **Love & Heroism:** Kasih yang melayani lebih dari standar (Magis) dan keberanian memimpin.
    - **Profil Lulusan:** 4C (Competence, Conscience, Compassion, Commitment) + 1L (Leadership).

10.  **Diskresi (Pembedaan Roh) & Spiritualitas:**
    - **Pembedaan Roh:** Membantu membedakan gerak roh baik (Consolation: damai, sukacita, kasih) dan roh jahat (Desolation: gelisah, takut, putus asa).
    - **AMDG (Ad Maiorem Dei Gloriam):** Segala sesuatu demi kemuliaan Tuhan yang lebih besar.
    - **Nilai St. Petrus Canisius:** Ketekunan, kesetiaan pada Gereja, dan edukasi yang membebaskan.

ATURAN OUTPUT / OUTPUT RULES:
- {language_instruction}
- Setiap saran atau materi pembelajaran harus memiliki "kedalaman" (menyentuh aspek Conscience/Compassion), bukan hanya teknis.
- Matematika wajib menggunakan LaTeX ($...$).
- Jika membuat TABEL, gunakan format Markdown standard.
- Berikan output yang TERSTRUKTUR rapi.
- FOKUS pada teks dan konten materi. Jangan menyertakan tag gambar.
"""

# --- 5. ENGINE: AUTO-DISCOVERY & SELF HEALING ---
class AIProvider:
    def __init__(self, api_key):
        self.api_key = api_key
        self.provider_name = "None"
        self.client = None
        self.available_models = []
        self.active_model = None
        self.is_valid = False
        
        if not api_key: return

        if api_key.startswith("gsk_"):
            self.provider_name = "Groq"
            self._setup_groq()
        else:
            self.provider_name = "Google"
            self._setup_google()
        
    def _setup_groq(self):
        try:
            self.client = Groq(api_key=self.api_key)
            models = self.client.models.list()
            self.available_models = [m.id for m in models.data if 'llama' in m.id or 'mixtral' in m.id]
            self.available_models.sort(key=lambda x: '70b' in x, reverse=True)
            if self.available_models:
                self.active_model = self.available_models[0]
                self.is_valid = True
        except: pass

    def _setup_google(self):
        try:
            genai.configure(api_key=self.api_key)
            priorities = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
            self.available_models = priorities 
            try:
                all_models = genai.list_models()
                real_models = [m.name.replace("models/", "") for m in all_models if 'generateContent' in m.supported_generation_methods]
                if real_models: self.available_models = [p for p in priorities if p in real_models] + [m for m in real_models if m not in priorities]
            except: pass
            
            self.active_model = self.available_models[0] if self.available_models else 'gemini-1.5-flash'
            self.is_valid = True
        except: pass

    def generate_stream(self, history, prompt, system_config, language_instruction, image_input=None, lib_text=""):
        if not self.is_valid: yield "⚠️ Error: API Key bermasalah."; return

        full_system = f"{get_system_prompt(language_instruction)}\n\n{system_config}"
        hist_str = "\n".join([f"{'USER' if m['role']=='user' else 'AI'}: {m['content']}" for m in history])
        final_prompt = f"HISTORY:\n{hist_str}\n\nSOURCES:\n{lib_text}\n\nUSER REQUEST:\n{prompt}"

        models_to_try = [self.active_model] + [m for m in self.available_models if m != self.active_model]
        success = False

        for model in models_to_try:
            if success: break
            try:
                if self.provider_name == "Google":
                    inputs = [f"SYSTEM_INSTRUCTION:\n{full_system}\n\nTASK:\n{final_prompt}"]
                    if image_input: inputs.append(image_input)
                    m = genai.GenerativeModel(model)
                    res = m.generate_content(inputs, stream=True)
                    for c in res: 
                        if c.text: yield c.text; success = True
                
                elif self.provider_name == "Groq":
                    if image_input: yield "ℹ️ [Groq: Gambar input diabaikan]\n"
                    stream = self.client.chat.completions.create(
                        messages=[{"role":"system","content":full_system},{"role":"user","content":final_prompt}],
                        model=model, stream=True
                    )
                    for c in stream:
                        txt = c.choices[0].delta.content
                        if txt: yield txt; success = True
            except: continue

# --- 6. LOGIC UI & HELPER (DOC ENGINE) ---
class DocEngine:
    @staticmethod
    def read(files):
        txt = ""; names = []
        for f in files:
            try:
                if f.name.endswith('.pdf'): txt += "".join([p.extract_text() for p in PyPDF2.PdfReader(f).pages])
                elif f.name.endswith('.docx'): txt += "\n".join([p.text for p in Document(f).paragraphs])
                elif f.name.endswith('.txt'): txt += f.getvalue().decode("utf-8")
                names.append(f.name)
            except: pass
        return txt, names

    @staticmethod
    def _set_table_borders(table):
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top'); top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
            left = OxmlElement('w:left'); left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '4')
            bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
            right = OxmlElement('w:right'); right.set(qn('w:val'), 'single'); right.set(qn('w:sz'), '4')
            tcBorders.append(top); tcBorders.append(left); tcBorders.append(bottom); tcBorders.append(right)
            tcPr.append(tcBorders)

    @staticmethod
    def _process_markdown_to_docx(doc, text):
        lines = text.split('\n')
        in_table = False
        table_data = []
        
        for line in lines:
            clean_line = line.strip()
            
            # Deteksi Tabel Markdown
            if "|" in clean_line and len(clean_line) > 2:
                if re.match(r'^\|?[\s-]+\|[\s-]+\|', clean_line): continue
                row_cells = [c.strip() for c in clean_line.split('|') if c.strip()]
                if row_cells:
                    if not in_table: in_table = True; table_data = [row_cells]
                    else: table_data.append(row_cells)
            else:
                if in_table:
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        for r_idx, row_content in enumerate(table_data):
                            for c_idx, cell_content in enumerate(row_content):
                                if c_idx < len(table.columns):
                                    cell = table.cell(r_idx, c_idx)
                                    cell.text = cell_content.replace('**', '')
                            DocEngine._set_table_borders(table)
                            doc.add_paragraph()
                        in_table = False; table_data = []

                # Heading & Text Formatting
                if clean_line.startswith('### '): doc.add_heading(clean_line.replace('### ', ''), level=3)
                elif clean_line.startswith('## '): doc.add_heading(clean_line.replace('## ', ''), level=2)
                elif clean_line.startswith('# '): doc.add_heading(clean_line.replace('# ', ''), level=1)
                elif clean_line:
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', clean_line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2]); run.bold = True
                        else: p.add_run(part.replace('$', ''))

    @staticmethod
    def create_word(history, title_text="Hasil Magis AI"):
        doc = Document()
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for msg in history:
            role_p = doc.add_heading(msg['role'].upper(), level=2)
            role_p.runs[0].font.color.rgb = RGBColor(27, 54, 93)
            DocEngine._process_markdown_to_docx(doc, msg['content'])
            doc.add_paragraph("-" * 20)

        bio = BytesIO(); doc.save(bio); return bio

# --- 7. APLIKASI UTAMA ---
if 'history' not in st.session_state: st.session_state.history = []
if 'library' not in st.session_state: st.session_state.library = {"text":"", "files":[]}

# API Handling
api_key = None
try:
    if "GOOGLE_API_KEY" in st.secrets: api_key = st.secrets["GOOGLE_API_KEY"]
    elif "GROQ_API_KEY" in st.secrets: api_key = st.secrets["GROQ_API_KEY"]
except: pass 

# --- SIDEBAR & SMART INPUT LOGIC ---
with st.sidebar:
    st.markdown("""
        <div style="text-align: center; margin-bottom: 20px;">
            <img src="https://i.imgur.com/UUCgyfV.png" width="110" style="filter: drop-shadow(0px 4px 6px rgba(0,0,0,0.1));">
        </div>
    """, unsafe_allow_html=True)
    
    # --- BAHASA / LANGUAGE SELECTOR ---
    selected_lang = st.selectbox("🌐 Bahasa / Language", ["Bahasa Indonesia", "English"])
    
    # Set Language Dictionary (T)
    if selected_lang == "Bahasa Indonesia":
        T = LANGUAGES["Bahasa Indonesia"]
    else:
        T = LANGUAGES["English"]
    
    # API Key Input
    if not api_key: 
        st.info(T["api_key_label"])
        api_key = st.text_input("API Key", type="password", label_visibility="collapsed")
        
    provider = AIProvider(api_key)
    if provider.is_valid: 
        st.markdown(f"<div class='status-ok'>{T['status_connected'].format(provider=provider.provider_name)}</div>", unsafe_allow_html=True)
    else: 
        st.markdown(f"<div class='status-err'>{T['status_waiting']}</div>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # MENU DENGAN IKON (Sekarang memiliki 4 opsi termasuk Chat Bebas)
    mode = st.selectbox(T["division_label"], T["divisions"])
    
    config_details = ""
    auto_prompt_template = "" 
    
    # Logic Pemilihan Mode
    
    # 1. AKADEMIK / ACADEMIC
    if mode == T["divisions"][0]: 
        st.markdown(T["academic_header"])
        with st.expander(T["class_expander"], expanded=True):
            input_kelas = st.selectbox(T["grade_label"], T["grades"])
            input_mapel = st.text_input(T["subject_label"], placeholder=T["subject_ph"])
            input_kd = st.text_area(T["topic_label"], placeholder=T["topic_ph"], height=80)
            
        with st.expander(T["params_expander"]):
            input_bloom = st.multiselect(T["bloom_label"], T["bloom_opts"], default=[T["bloom_opts"][3], T["bloom_opts"][4]])
            input_difficulty = st.select_slider(T["difficulty_label"], options=T["difficulty_opts"])
            
        with st.expander(T["style_expander"]):
            input_gaya = st.selectbox(T["style_label"], T["style_opts"])
            input_ipp_focus = st.multiselect(T["ipp_label"], ["Context", "Experience", "Reflection", "Action", "Evaluation"], default=["Reflection"])

        config_details = f"CONFIG: {input_kelas}, {input_mapel}, {input_gaya}, IPP {','.join(input_ipp_focus)}"
        
        # Template Prompt
        if selected_lang == "Bahasa Indonesia":
            auto_prompt_template = (
                f"Saya guru {input_mapel} untuk kelas {input_kelas}. \n"
                f"Topik: {input_kd if input_kd else '[Isi Topik]'}. \n\n"
                f"Tolong buatkan [Rencana Pembelajaran / 5 Soal PG / Studi Kasus] "
                f"dengan level kognitif {', '.join(input_bloom)} dan tingkat kesulitan {input_difficulty}. "
                f"Tekankan pada aspek {', '.join(input_ipp_focus)}."
            )
        else:
            auto_prompt_template = (
                f"I am a {input_mapel} teacher for {input_kelas}. \n"
                f"Topic: {input_kd if input_kd else '[Enter Topic]'}. \n\n"
                f"Please create [Lesson Plan / 5 Multiple Choice Questions / Case Study] "
                f"with cognitive levels {', '.join(input_bloom)} and difficulty level {input_difficulty}. "
                f"Emphasize the {', '.join(input_ipp_focus)} aspects."
            )

    # 2. PASTORAL
    elif mode == T["divisions"][1]:
        st.markdown(T["pastoral_header"])
        with st.expander(T["counsel_expander"], expanded=True):
            pas_subjek = st.selectbox(T["subject_label"], T["subjects"])
            pilihan_masalah = st.selectbox(T["issue_label"], T["issue_opts"])
            if "Lain" in pilihan_masalah or "Other" in pilihan_masalah:
                pas_masalah = st.text_input(T["issue_label"], placeholder=T["issue_custom_ph"])
            else:
                pas_masalah = pilihan_masalah
            pas_metode = st.radio(T["method_label"], T["methods"])
        
        config_details = f"PASTORAL CONFIG: Subject {pas_subjek}, Issue {pas_masalah}, Method {pas_metode}"
        
        if selected_lang == "Bahasa Indonesia":
            auto_prompt_template = (
                f"Saya sedang mendampingi seorang {pas_subjek} yang mengalami pergumulan tentang {pas_masalah if pas_masalah else '[Isi Masalah]'}. \n\n"
                f"Mohon berikan panduan percakapan atau refleksi menggunakan pendekatan {pas_metode}. "
                f"Tujuannya adalah membantu subjek menemukan kedamaian (konsolasi) dan mengambil keputusan yang tepat."
            )
        else:
            auto_prompt_template = (
                f"I am counseling a {pas_subjek} who is struggling with {pas_masalah if pas_masalah else '[Enter Issue]'}. \n\n"
                f"Please provide a conversation guide or reflection using the {pas_metode} approach. "
                f"The goal is to help them find peace (consolation) and make the right decision."
            )
        
    # 3. MANAJEMEN / MANAGEMENT
    elif mode == T["divisions"][2]: 
        st.markdown(T["management_header"])
        man_jenis = st.selectbox(T["doc_type_label"], T["doc_types"])
        man_tone = st.select_slider(T["tone_label"], options=T["tones"])
        man_topik = st.text_input(T["topic_mgmt_label"], placeholder=T["topic_mgmt_ph"])
        
        config_details = f"MANAGEMENT CONFIG: Doc {man_jenis}, Tone {man_tone}"
        
        if selected_lang == "Bahasa Indonesia":
            auto_prompt_template = (
                f"Buatkan draf {man_jenis} bertema '{man_topik if man_topik else '[Isi Topik]'}'. \n\n"
                f"Gunakan nada bicara yang {man_tone}. "
                f"Pastikan struktur dokumen rapi dan sesuai standar institusi pendidikan Jesuit."
            )
        else:
            auto_prompt_template = (
                f"Draft a {man_jenis} regarding '{man_topik if man_topik else '[Enter Topic]'}.' \n\n"
                f"Use a {man_tone} tone. "
                f"Ensure the document structure is neat and follows Jesuit educational institution standards."
            )
            
    # 4. CHAT BEBAS (GENERAL) - DIKEMBALIKAN KE SISTEM
    else:
        st.markdown(T["general_header"])
        st.info(T["general_help"])
        config_details = "MODE: General Discussion / Free Chat. Context: Jesuit Education Companion."
        auto_prompt_template = "" # Kosongkan agar user bebas mengetik

    # --- AKSESIBILITAS / TAMPILAN (FITUR BARU) ---
    st.markdown("---")
    with st.expander(T["access_header"], expanded=False):
        font_size = st.slider(T["font_size_label"], 14, 28, 16)
        
    # Inject CSS Dinamis berdasarkan Font Size
    inject_css(font_size)

    # FOOTER AREA (FIXED WITH CREDIT)
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button(T["reset_btn"], use_container_width=True): 
        st.session_state.history = []
        st.rerun()
        
    # Footer dengan Credit dan Detail
    st.markdown(f"""
        <div class="sidebar-footer">
            <strong>Magis AI v7.3</strong><br>
            {T["footer_design"]}<br>
            Kolese Kanisius Jakarta | 2026
        </div>
    """, unsafe_allow_html=True)

# --- MAIN UI ---
c1,c2 = st.columns([3,1])
with c1: 
    st.markdown(f'''
    <div class="magis-title">{T["magis_title"]}</div>
    <div class="magis-tagline">{T["magis_tagline"]}</div>
    <div class="magis-badge">{T["mode_active"]}: {mode}</div>
    ''', unsafe_allow_html=True)

with st.expander(T["upload_expander"], expanded=False):
    st.markdown(T["upload_help"])
    files = st.file_uploader(T["file_label"], accept_multiple_files=True)
    img_up = st.file_uploader(T["img_label"], type=['png','jpg'])
    if files:
        t, n = DocEngine.read(files)
        st.session_state.library = {"text": t, "files": n}
        st.success(T["success_msg"].format(count=len(n)))

# --- CHAT DISPLAY ---
chat_container = st.container()
with chat_container:
    for m in st.session_state.history:
        st.markdown(f"<div class='{'bubble-user' if m['role']=='user' else 'bubble-ai'}'>{m['content'].replace('[DOC_CONTEXT]','')}</div>", unsafe_allow_html=True)
    st.markdown("<div style='height: 50px;'></div>", unsafe_allow_html=True)

# --- SMART INPUT AREA ---
st.markdown("---")
st.markdown(T["workspace_header"])

with st.form(key='smart_input_form', clear_on_submit=True):
    # Menggunakan hash prompt untuk key agar input field ter-refresh saat mode ganti
    prompt_key = f"input_{hash(auto_prompt_template) if auto_prompt_template else 'free'}" 
    
    user_in = st.text_area(
        T["instruction_label"], 
        value=auto_prompt_template, 
        height=250, 
        key=prompt_key
    )
    
    col_act1, col_act2 = st.columns([1, 5])
    with col_act1:
        submitted = st.form_submit_button(T["submit_btn"], use_container_width=True)
    with col_act2:
        st.caption(T["tip_msg"])

# LOGIC PEMROSESAN
if submitted and user_in and provider.is_valid:
    final_msg = user_in
    if st.session_state.library["text"]: final_msg += " [DOC_CONTEXT]"
    st.session_state.history.append({"role":"user", "content":final_msg})
    st.rerun()

if st.session_state.history and st.session_state.history[-1]['role'] == 'user':
    with st.spinner(T["spinner_msg"]):
        full_res = ""
        box = st.empty()
        last_usr = st.session_state.history[-1]['content']
        img_data = PIL.Image.open(img_up) if img_up else None
        
        # Kirim System Instruction sesuai Bahasa & Config
        for chk in provider.generate_stream(st.session_state.history[:-1], last_usr, config_details, T["system_lang_instruction"], img_data, st.session_state.library["text"]):
            full_res += chk
            box.markdown(f"<div class='bubble-ai'>{full_res}</div>", unsafe_allow_html=True)
            
        st.session_state.history.append({"role":"assistant", "content":full_res})
        st.rerun()

if st.session_state.history:
    st.markdown(T["export_header"])
    docx_title = "Hasil-MagisAI.docx" if selected_lang == "Bahasa Indonesia" else "MagisAI-Result.docx"
    doc_header = "Hasil Magis AI" if selected_lang == "Bahasa Indonesia" else "Magis AI Result"
    
    docx = DocEngine.create_word(st.session_state.history, title_text=doc_header)
    st.download_button(
        label=T["download_label"], 
        data=docx, 
        file_name=docx_title, 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
